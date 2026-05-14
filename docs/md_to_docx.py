"""Convert highres_report.md into highres_report.docx.

Minimal Markdown subset: # / ## / ### headings, paragraphs, **bold**, *italic*,
inline `code`, fenced ```code blocks```, GFM-style tables, image lines as captions.
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor


INLINE_PATTERNS = [
    (re.compile(r"\*\*(.+?)\*\*"), "bold"),
    (re.compile(r"`([^`]+)`"), "code"),
    (re.compile(r"\*(.+?)\*"), "italic"),
]


def add_runs(paragraph, text):
    """Parse a line for **bold**, `code`, *italic* and emit styled runs."""
    cursor = 0
    n = len(text)
    while cursor < n:
        best = None
        for pat, kind in INLINE_PATTERNS:
            m = pat.search(text, cursor)
            if m and (best is None or m.start() < best[0].start()):
                best = (m, kind)
        if not best:
            paragraph.add_run(text[cursor:])
            return
        m, kind = best
        if m.start() > cursor:
            paragraph.add_run(text[cursor:m.start()])
        run = paragraph.add_run(m.group(1))
        if kind == "bold":
            run.bold = True
        elif kind == "italic":
            run.italic = True
        elif kind == "code":
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        cursor = m.end()


def style_caption(doc, caption_text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"[Figure placeholder] {caption_text}")
    r.italic = True
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    r.font.size = Pt(10)


def emit_table(doc, rows):
    if not rows:
        return
    # Filter empty cells from each row trailing
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Light Grid Accent 1"
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            tc = table.cell(i, j)
            tc.text = ""
            p = tc.paragraphs[0]
            add_runs(p, cell)
            for run in p.runs:
                run.font.size = Pt(9)
                if i == 0:
                    run.bold = True
            tc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def parse_md(md_text):
    lines = md_text.splitlines()
    doc = Document()

    # Set default font to something pleasant for body text
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Wider margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    i = 0
    while i < len(lines):
        line = lines[i]

        # Code block (fenced)
        if line.startswith("```"):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # consume closing fence
            p = doc.add_paragraph()
            r = p.add_run("\n".join(code_lines))
            r.font.name = "Consolas"
            r.font.size = Pt(9.5)
            continue

        # Heading
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            h = doc.add_heading(line[2:].strip(), level=0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Horizontal rule
        elif line.strip() == "---":
            doc.add_paragraph()
        # Table
        elif line.lstrip().startswith("|") and "|" in line:
            # Collect rows
            rows = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not all(set(c) <= set("-:") for c in cells):  # skip separator row
                    rows.append(cells)
                i += 1
            emit_table(doc, rows)
            continue
        # Image (caption only, since we don't have the file)
        elif line.startswith("!["):
            m = re.match(r"!\[(.*?)\]\((.*?)\)", line)
            if m:
                style_caption(doc, m.group(1))
            else:
                doc.add_paragraph(line)
        # Italic-only paragraph (e.g. byline)
        elif line.startswith("*") and line.rstrip().endswith("*") and not line.startswith("**"):
            p = doc.add_paragraph()
            r = p.add_run(line.strip("*").strip())
            r.italic = True
        # Numbered list item
        elif re.match(r"^\d+\.\s", line):
            p = doc.add_paragraph(style="List Number")
            text = re.sub(r"^\d+\.\s+", "", line)
            add_runs(p, text)
        # Plain paragraph (skip blank lines that are just structure)
        elif line.strip():
            p = doc.add_paragraph()
            add_runs(p, line)
        else:
            # Blank line — emit a paragraph break implicitly
            pass

        i += 1
    return doc


def main():
    src = Path(sys.argv[1] if len(sys.argv) > 1 else "highres_report.md")
    dst = src.with_suffix(".docx")
    md_text = src.read_text(encoding="utf-8")
    doc = parse_md(md_text)
    doc.save(dst)
    print(f"wrote {dst} ({dst.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()
